import streamlit as st
import requests
import io
from docx import Document
from fpdf import FPDF

import sounddevice as sd
import numpy as np
import wave
import time
import os

print(os.getcwd())
print(os.listdir())

# Initialize session state for storing transcription
if 'transcription' not in st.session_state:
    st.session_state.transcription = None
if 'current_file' not in st.session_state:
    st.session_state.current_file = None
if 'duration' not in st.session_state:
    st.session_state.duration = 5

# Function to record audio
def record_audio(duration=st.session_state.duration, samplerate=44100):
    st.info(f"Recording for {duration} seconds...")

    # Create a progress bar and status text
    progress_bar = st.progress(0)
    status_text = st.empty()

    audio_data = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=1, dtype=np.int16)

    # Update progress while recording
    start_time = time.time()
    while (time.time() - start_time) < duration:
        elapsed = time.time() - start_time
        remaining = max(0, duration - elapsed)
        progress = min(1.0, elapsed / duration)

        progress_bar.progress(progress)
        status_text.text(f"Time remaining: {remaining:.1f} seconds")
        time.sleep(0.1)  # Update 10 times per second

    # Ensure recording is complete
    sd.wait()

    # Final update
    progress_bar.progress(1.0)
    status_text.text("Recording complete!")

    # Save to a BytesIO object
    audio_io = io.BytesIO()
    with wave.open(audio_io, 'wb') as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(samplerate)
        wf.writeframes(audio_data.tobytes())
    audio_io.seek(0)

    wav_data = audio_io.read()

    # Create a new BytesIO object that acts like an uploaded file
    file_like = io.BytesIO(wav_data)
    file_like.name = "recording.wav"  # Important: give it a name with .wav extension

    return file_like


# URL of your backend on laptop
# BACKEND_URL = "http://127.0.0.1:8000/upload"


# production link
BACKEND_URL = "https://transcriptai.onrender.com/upload"

# Title of the app
st.title("Audio Transcription + Document Generation by Cesar")

# Custom CSS for button styling
st.markdown("""
<style>
/* Center the Transcribe button */
    div.stButton > button {
        display: block;
        margin: 0 auto;
    }
    /* Target the first download button (Word) */
    div.stDownloadButton:nth-of-type(1) button {
        background-color: #1E90FF !important;
        color: white !important;
        border: none !important;
    }
    div.stDownloadButton:nth-of-type(1) button:hover {
        background-color: #0066CC !important;
    }

    /* Target the second download button (PDF) */
    div.stDownloadButton:nth-of-type(2) button {
        background-color: #FF4444 !important;
        color: white !important;
        border: none !important;
    }
    div.stDownloadButton:nth-of-type(2) button:hover {
        background-color: #CC0000 !important;
    }

    /* Center buttons in their columns */
    .stDownloadButton {
        display: flex;
        justify-content: center;
    }

    /* Custom button styles for links */
    .link-button {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 10px;
        padding: 10px;
        border-radius: 8px;
        color: white;
        text-decoration: none;
        font-weight: bold;
        font-size: 16px;
        width: 100%;
        text-align: center;
    }
    .linkedin { background-color: #9adaf5; }
    .github { background-color: #f59453; }
    .portfolio { background-color: #8b00c7; }
    .link-button:hover { opacity: 0.8; }
    .icon { width: 24px; height: 24px; }
    .link-button:active { transform: scale(0.98); } /* Slight click effect */
</style>
""", unsafe_allow_html=True)

# Upload the audio file
uploaded_file = st.file_uploader("Upload an audio file", type=["mp3", "wav", "m4a"])

# Clear transcription if a new file is uploaded
if uploaded_file and uploaded_file != st.session_state.current_file:
    st.session_state.transcription = None
    st.session_state.current_file = uploaded_file

# If file is uploaded, show the player and send to backend
if uploaded_file:
    st.audio(uploaded_file, format="audio/wav")
    st.session_state.current_file = uploaded_file

# st.session_state.duration = st.slider("Or select recording duration (seconds) and start recording", min_value=1, max_value=30, value=5)
#
# if st.button("Record Audio"):
#     recorded_audio = record_audio()
#
#     if recorded_audio:  # Only proceed if recording was successful
#         # Create a copy of the audio data for playback
#         playback_audio = io.BytesIO(recorded_audio.getvalue())
#         st.audio(playback_audio, format="audio/wav")
#
#         # Reset and store the original for transcription
#         recorded_audio.seek(0)
#         st.session_state.current_file = recorded_audio
#
#         # Clear any previous transcription
#         st.session_state.transcription = None
#         st.success("Audio recorded successfully! Click 'Transcribe' to process.")
#     else:
#         st.error("Recording failed. Please try again.")

if st.session_state.current_file and st.button("Transcribe"):
    # Send audio file to backend for transcription
    files = {"file": st.session_state.current_file}
    response = requests.post(BACKEND_URL, files=files, timeout=120)

    if response.status_code == 200:
        # Convert transcription to Word document
        st.session_state.transcription = response.json().get("transcript")
    else:
        st.error("Error during transcription.")


# Function to generate Word document
def generate_word():
    # Convert transcription to Word document
    doc = Document()
    doc.add_heading(f"{st.session_state.current_file.name} Transcription", 0)
    doc.add_paragraph(st.session_state.transcription)

    # Save Word file to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# Function to generate PDF document
# def generate_pdf():
#     # Get the current working directory
#     current_dir = os.path.dirname(os.path.abspath(__file__))
#
#     # Construct the path to the font file relative to the current directory
#     font_path = os.path.join(current_dir, 'DejaVuSansCondensed.ttf')
#
#     # Create PDF document
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.add_font("DejaVu", "", font_path, uni=True)
#     pdf.set_font("DejaVu", size=20)
#     pdf.cell(200, 10, txt=f"{st.session_state.current_file.name} Transcription", ln=True, align='C')
#
#     pdf.set_font("DejaVu", size=12)
#     pdf.cell(200, 10, txt="", ln=True, align='C')
#     pdf.multi_cell(0, 10, st.session_state.transcription)
#
#     # Save PDF to BytesIO
#     pdf_output = pdf.output(dest="S").encode("latin1")  # Convert to bytes compatible with PDF
#     pdf_io = io.BytesIO(pdf_output)
#     pdf_io.seek(0)
#     return pdf_io


# Display download buttons only if transcription exists in session state
# if st.session_state.transcription and st.session_state.current_file:
#     st.success("Transcription completed!")
#
#     # Create centered columns for download buttons
#     col1, col2 = st.columns(2)
#
#     with col1:
#         # Download Word button
#         doc_io = generate_word()  # Generate Word file
#         st.download_button(
#             label="📝 Download Word",
#             data=doc_io,
#             file_name=f"{st.session_state.current_file.name}_transcription.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             use_container_width=True
#         )
#
#     with col2:
#         # Download PDF button
#         pdf_io = generate_pdf()  # Generate PDF file
#         st.download_button(
#             label="📄 Download PDF",
#             data=pdf_io,
#             file_name=f"{st.session_state.current_file.name}_transcription.pdf",
#             mime="application/pdf",
#             use_container_width=True
#         )

# Display download buttons only if transcription exists in session state


if st.session_state.transcription and st.session_state.current_file:
    st.success("Transcription completed!")

    # Download Word button
    doc_io = generate_word()  # Generate Word file
    st.download_button(
            label="📝 Download Word",
            data=doc_io,
            file_name=f"{st.session_state.current_file.name}_transcription.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
    )


# Adaugă spațiu înainte de secțiunea de contact
st.markdown("<br><br>", unsafe_allow_html=True)

# Afișează un mesaj de avertizare
st.warning("This is a beta version of this app. Feel free to reach out with suggestions or questions.")

# Profile picture URL (replace with actual image URL)
profile_image_url = "https://leoni-cesar-portfolio.netlify.app/assets/cesar.webp"

# Create two columns for the profile picture and links
col1, col2 = st.columns([1, 2])  # Adjust width ratio as needed


with col1:
    st.markdown("<div style='display: flex; justify-content: center; align-items: center; height: 100%;'>", unsafe_allow_html=True)
    st.image(profile_image_url, width=190, caption="")
    st.markdown("</div>", unsafe_allow_html=True)


with col2:
    st.markdown("<div style='display: flex; justify-content: center; align-items: center; height: 100%;'>", unsafe_allow_html=True)
    st.markdown(
        """
        <style>
            .link-button {
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 10px;
                padding: 10px;
                border-radius: 8px;
                color: black;
                text-decoration: none;
                font-weight: bold;
                font-size: 16px;
                width: 100%;
                text-align: center;
                cursor: pointer;
            }

            .linkedin { background-color: #9adaf5; color: blue; }
            .github { background-color: #f59453; }
            .portfolio { background-color: #8b00c7; color: yellow; }

            .link-button:hover { opacity: 0.8; }

            .icon { width: 24px; height: 24px; }
        </style>

        <a class="link-button linkedin" href="https://www.linkedin.com/in/cesarleoni29/" target="_blank">
            <img class="icon" src="https://cdn-icons-png.flaticon.com/512/174/174857.png"> LinkedIn
        </a>

        <br>

        <a class="link-button github" href="https://github.com/CesarLeoni" target="_blank">
            <img class="icon" src="https://cdn-icons-png.flaticon.com/512/25/25231.png"> GitHub
        </a>

        <br>

        <a class="link-button portfolio" href="https://leoni-cesar-portfolio.netlify.app/" target="_blank">
            <img class="icon" src="https://cdn-icons-png.flaticon.com/512/616/616489.png"> Portfolio
        </a>
        """,
        unsafe_allow_html=True
    )

    st.markdown("</div>", unsafe_allow_html=True)

